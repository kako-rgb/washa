import os
import sys
import json
from docx import Document

def extract_text_from_docx(docx_path):
    try:
        # Check if file exists
        if not os.path.exists(docx_path):
            print(f"Error: File '{docx_path}' not found.")
            return None
            
        # Load the document
        doc = Document(docx_path)
        
        # Extract data
        data = {
            'paragraphs': [],
            'tables': []
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only include non-empty paragraphs
                data['paragraphs'].append({
                    'text': para.text,
                    'length': len(para.text)
                })
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            
            data['tables'].append({
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.rows[0].cells) if table.rows else 0,
                'data': table_data
            })
            
        return data
        
    except Exception as e:
        print(f"Error extracting data from DOCX: {str(e)}")
        return None

def save_to_json(data, output_path):
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f'Data saved to {output_path}')
        return True
    except Exception as e:
        print(f"Error saving JSON: {str(e)}")
        return False

def main():
    # Check if we have the right number of arguments
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <docx_file_path> [output_json_path]")
        return False
    
    # Get the DOCX file path
    docx_path = sys.argv[1]
    
    # Get the output JSON path (optional)
    if len(sys.argv) >= 3:
        output_path = sys.argv[2]
    else:
        # Default output path is the same as input but with .json extension
        output_path = os.path.splitext(docx_path)[0] + '.json'
    
    # Extract data from DOCX
    data = extract_text_from_docx(docx_path)
    
    if data:
        # Save to JSON
        success = save_to_json(data, output_path)
        if success:
            print("Data extraction completed successfully.")
            return True
    
    print("Failed to extract data from the document.")
    return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
