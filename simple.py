from docx import Document
print("Starting script...")
doc = Document("olddata/may.docx")
print(f"Document has {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
