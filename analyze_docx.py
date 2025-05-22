from docx import Document
import json
import os

def extract_docx_data(docx_path):
    """Extract data from a DOCX file and return structured information."""
    print(f'Opening document: {docx_path}')
    doc = Document(docx_path)
    
    # Basic document stats
    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'total_tables': len(doc.tables),
        'non_empty_paragraphs': sum(1 for p in doc.paragraphs if p.text.strip()),
    }
    
    # Extract paragraph data
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            paragraphs.append({
                'index': i,
                'text': para.text,
                'length': len(para.text)
            })
    
    # Extract table data
    tables = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0
        
        # Extract table content
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        tables.append({
            'index': i,
            'rows': rows,
            'columns': cols,
            'data': table_data
        })
    
    return {
        'stats': stats,
        'paragraphs': paragraphs[:100],  # Limit to first 100 paragraphs for brevity
        'tables': tables[:20]  # Limit to first 20 tables for brevity
    }

def save_to_json(data, output_path):
    """Save extracted data to a JSON file."""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f'Data saved to {output_path}')

def print_summary(data):
    """Print a summary of the extracted data."""
    stats = data['stats']
    print('\n=== DOCUMENT SUMMARY ===')
    print(f'Total paragraphs: {stats["total_paragraphs"]}')
    print(f'Non-empty paragraphs: {stats["non_empty_paragraphs"]}')
    print(f'Total tables: {stats["total_tables"]}')
    
    # Print sample paragraphs
    print('\n=== SAMPLE PARAGRAPHS ===')
    for i, para in enumerate(data['paragraphs'][:5]):
        text = para['text']
        print(f'Paragraph {para["index"]+1}: {text[:100]}...' if len(text) > 100 else f'Paragraph {para["index"]+1}: {text}')
    
    # Print sample tables
    print('\n=== SAMPLE TABLES ===')
    for i, table in enumerate(data['tables'][:3]):
        print(f'Table {table["index"]+1}: {table["rows"]} rows x {table["columns"]} columns')
        
        # Print first few rows of the table
        if table['rows'] > 0:
            print('  Sample data:')
            for j, row in enumerate(table['data'][:3]):
                print(f'  Row {j+1}: {row[:3]}' + ('...' if len(row) > 3 else ''))

def main():
    docx_path = 'olddata/may.docx'
    output_path = 'olddata/may_data.json'
    
    try:
        # Extract data from DOCX
        data = extract_docx_data(docx_path)
        
        # Print summary
        print_summary(data)
        
        # Save to JSON
        save_to_json(data, output_path)
        
        print('\nData extraction completed successfully.')
        print(f'Full data saved to {output_path}')
        print('You can now use this data for further processing.')
        
    except Exception as e:
        print(f'Error: {str(e)}')

if __name__ == '__main__':
    main()